#!/usr/bin/env python3
"""Generate A4 sell-sheet Word (.docx) for TailFin Marine."""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG = os.path.join(ROOT, "assets", "img")
OUT = os.path.join(ROOT, "TailFin_Marine_SellSheet.docx")

NAVY = RGBColor(0x0f, 0x27, 0x40)
TEAL = RGBColor(0x1b, 0xa3, 0x9c)
AMBER = RGBColor(0xf5, 0xa6, 0x23)
GRAY = RGBColor(0x5b, 0x6b, 0x78)

doc = Document()

# ---- A4 page setup ----
sec = doc.sections[0]
sec.page_width = Mm(210); sec.page_height = Mm(297)
sec.top_margin = Mm(9); sec.bottom_margin = Mm(9)
sec.left_margin = Mm(9); sec.right_margin = Mm(9)

# default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, v in (("top",top),("bottom",bottom),("start",left),("end",right)):
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:w"), str(v)); e.set(qn("w:type"), "dxa"); m.append(e)
    tcPr.append(m)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"),"none"); e.set(qn("w:sz"),"0")
        borders.append(e)
    tblPr.append(borders)

def para(cell, text, size=9, bold=False, color=None, align=None, space_after=2):
    p = cell.paragraphs[0] if (cell.paragraphs and not cell.paragraphs[0].runs) else cell.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return p

# ===== HEADER =====
hdr = doc.add_table(rows=1, cols=1); no_borders(hdr)
c = hdr.rows[0].cells[0]
shade(c, "0F2740"); set_cell_margins(c, 120, 120, 160, 160)
p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
r = p.add_run("TailFin Marine"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0xff,0xff,0xff)
p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(0)
r2 = p2.add_run("Modular Off-Grid Fish-Processing Systems · 模块化离网钓后处理系统")
r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0xb9,0xc9,0xd6)
p3 = c.add_paragraph(); p3.paragraph_format.space_before = Pt(3)
r3 = p3.add_run("CATCH. CLEAN. COOK. ANYWHERE."); r3.bold = True; r3.font.size = Pt(13); r3.font.color.rgb = AMBER
p4 = c.add_paragraph(); p4.paragraph_format.space_after = Pt(0)
r4 = p4.add_run("Engineered for serious anglers, everywhere. The world's first powered, modular catch-to-freezer workstation.  ·  PRODUCT SELL SHEET 招商简版")
r4.font.size = Pt(8); r4.font.color.rgb = RGBColor(0xcd,0xd9,0xe2)

# ===== LEAD: hero + USP =====
lead = doc.add_table(rows=1, cols=2); no_borders(lead)
lead.columns[0].width = Mm(88); lead.columns[1].width = Mm(94)
lc = lead.rows[0].cells[0]; set_cell_margins(lc, 60, 60, 60, 60)
shade(lc, "F4F7F9")
ip = lc.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
ip.add_run().add_picture(os.path.join(IMG, "pro-product.png"), width=Mm(72))
cap = lc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cap.add_run("Mobile Fish Factory (Pro) — the complete powered system")
cr.font.size = Pt(7.5); cr.font.color.rgb = GRAY

rc = lead.rows[0].cells[1]; set_cell_margins(rc, 60, 60, 60, 60)
h = rc.paragraphs[0]; h.paragraph_format.space_after = Pt(3)
hr = h.add_run("Why TailFin wins the dock"); hr.bold = True; hr.font.size = Pt(10.5); hr.font.color.rgb = NAVY
for n, t in [
    ("1","Powered, not passive. Built-in 12/24V electric descaler + vacuum sealer clean & preserve a full catch in minutes — no shore power."),
    ("2","Fully modular & off-grid. Battery, scaler, fillet knife, vacuum cartridge and hose swap in seconds; runs on 12/24V or its own battery."),
    ("3","Dock to freezer, one station. Scale, fillet and vacuum-seal at the tailgate, deck or ice hole — 12 recurring consumables drive repeat revenue."),
]:
    pp = rc.add_paragraph(); pp.paragraph_format.space_after = Pt(3)
    rn = pp.add_run(f"{n}. "); rn.bold = True; rn.font.color.rgb = TEAL; rn.font.size = Pt(9)
    rt = pp.add_run(t); rt.font.size = Pt(8.5)

# ===== LINEUP =====
doc.add_paragraph().add_run("The Lineup · 产品矩阵").bold = True
lineup = doc.add_table(rows=2, cols=4); no_borders(lineup)
for col in lineup.columns:
    col.width = Mm(45.5)
products = [
    ("mini-product.png", "The Pocket Scaler (Mini)", "掌上电动去鳞器", "Palm-sized 12V descaler for kayak, travel & quick jobs.", "$149"),
    ("standard-product.png", "Tailgate Workstation", "标准模块化工作站", "The core modular station for truck tailgate & boat deck.", "$499"),
    ("pro-product.png", "Mobile Fish Factory (Pro)", "移动鱼类加工站", "Full powered system: descaler + vacuum + battery + knife.", "$999"),
    ("acc-kit.png", "Pro Upgrade Kit", "专业升级套件", "Turn any Standard into a Pro: vacuum + battery + knife.", "$499"),
]
for i, (img, name, zh, desc, price) in enumerate(products):
    imgcell = lineup.rows[0].cells[i]
    shade(imgcell, "F4F7F9"); set_cell_margins(imgcell, 40, 40, 40, 40)
    ip = imgcell.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(os.path.join(IMG, img), width=Mm(34))
    body = lineup.rows[1].cells[i]; set_cell_margins(body, 50, 50, 50, 50)
    shade(body, "FFFFFF")
    para(body, name, size=8.5, bold=True, color=NAVY, space_after=0)
    para(body, zh, size=7.5, color=GRAY, space_after=1)
    para(body, desc, size=7, color=GRAY, space_after=2)
    para(body, price, size=10, bold=True, color=AMBER, space_after=0)

# ===== PARTNER =====
doc.add_paragraph().add_run("Why distributors partner with us · 分销合作理由").bold = True
pr = doc.add_table(rows=1, cols=2); no_borders(pr)
pr.columns[0].width = Mm(125); pr.columns[1].width = Mm(57)
lc2 = pr.rows[0].cells[0]; shade(lc2, "0F2740"); set_cell_margins(lc2, 90, 90, 120, 120)
points = [
    ("Category gap:", "rivals sell passive HDPE boards ($66–$120); we own the powered niche."),
    ("Healthy margins:", "up to 50% off MSRP for national distributors."),
    ("Recurring revenue:", "12 consumables per catch cycle — blades $29, bags $19, battery $149…"),
    ("Full support:", "marketing kit, MAP & region protection, certifications & after-sales."),
]
for head, body in points:
    pp = lc2.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
    rh = pp.add_run(head + " "); rh.bold = True; rh.font.size = Pt(8); rh.font.color.rgb = RGBColor(0xff,0xff,0xff)
    rb = pp.add_run(body); rb.font.size = Pt(8); rb.font.color.rgb = RGBColor(0xea,0xf1,0xf6)
rc2 = pr.rows[0].cells[1]; shade(rc2, "FFFFFF"); set_cell_margins(rc2, 80, 80, 100, 100)
para(rc2, "Distributor Snapshot · 合作速览", size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
snap = [("MOQ","200 units / container"),("National dist.","50% off MSRP"),("Retail MAP","Protected"),("Rebate","Tiered, 3–8%"),("Lead time","30–45 days")]
for k, v in snap:
    sp = rc2.add_paragraph(); sp.paragraph_format.space_after = Pt(1)
    rk = sp.add_run(f"{k}: "); rk.font.size = Pt(7.5); rk.font.color.rgb = NAVY
    rv = sp.add_run(v); rv.bold = True; rv.font.size = Pt(7.5); rv.font.color.rgb = AMBER

# ===== FOOTER =====
foot = doc.add_table(rows=1, cols=1); no_borders(foot)
fc = foot.rows[0].cells[0]; shade(fc, "0A1B2C"); set_cell_margins(fc, 90, 90, 140, 140)
fp = fc.paragraphs[0]; fp.paragraph_format.space_after = Pt(0)
fr = fp.add_run("Partner with TailFin Marine · 期待与全球分销商、户外连锁及文旅渠道合作   |   ")
fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0x9f,0xb3,0xc2)
fr2 = fp.add_run("hello@tailfinmarine.com · www.tailfinmarine.com · Booth & full catalog on request")
fr2.font.size = Pt(8); fr2.bold = True; fr2.font.color.rgb = AMBER

doc.save(OUT)
print("saved", OUT, os.path.getsize(OUT), "bytes")
