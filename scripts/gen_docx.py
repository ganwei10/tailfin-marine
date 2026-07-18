# -*- coding: utf-8 -*-
"""Generate business-grade .docx (native Word tables) for TailFin Marine, ZH & EN."""
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0f, 0x1e, 0x37)
TEAL = RGBColor(0x0f, 0x6e, 0x56)
GRAY = RGBColor(0x5b, 0x66, 0x75)
TEALBG = "E6F3EF"
GRAYBG = "F2F4F7"


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def set_cjk(run, font='Microsoft YaHei'):
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)


def add_title(doc, text, en_font):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    set_cjk(run, en_font)
    return p


def add_h2(doc, text, en_font):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL
    set_cjk(run, en_font)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '0F6E56')
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def add_h3(doc, text, en_font):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL
    set_cjk(run, en_font)
    return p


def add_para(doc, text, en_font, muted=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GRAY if muted else RGBColor(0x1d, 0x27, 0x33)
    set_cjk(run, en_font)
    return p


def add_bullets(doc, items, en_font):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(it)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1d, 0x27, 0x33)
        set_cjk(run, en_font)


def add_table(doc, headers, rows, en_font):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], TEALBG)
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = NAVY
        set_cjk(run, en_font)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1:
                set_cell_bg(cells[ci], GRAYBG)
            para = cells[ci].paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1d, 0x27, 0x33)
            set_cjk(run, en_font)
    return t


def build(lang):
    zh = (lang == 'zh')
    font = 'Microsoft YaHei' if zh else 'Calibri'
    doc = Document()
    # base style font
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = font
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    T = TXT[lang]
    add_title(doc, "TailFin Marine", font)
    sub = doc.add_paragraph(); r = sub.add_run(T['sub']); r.font.size = Pt(11); r.bold = True; r.font.color.rgb = NAVY; set_cjk(r, font)
    add_para(doc, T['lede'], font, muted=True)

    for sec in T['sections']:
        add_h2(doc, sec['h'], font)
        for block in sec['body']:
            kind = block[0]
            if kind == 'p':
                add_para(doc, block[1], font)
            elif kind == 'muted':
                add_para(doc, block[1], font, muted=True, size=9.5)
            elif kind == 'h3':
                add_h3(doc, block[1], font)
            elif kind == 'ul':
                add_bullets(doc, block[1], font)
            elif kind == 'table':
                add_table(doc, block[1], block[2], font)
            elif kind == 'callout':
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), TEALBG)
                pPr.append(shd)
                run = p.add_run(block[1]); run.font.size = Pt(10.5); run.bold = True; run.font.color.rgb = TEAL; set_cjk(run, font)

    out = 'TailFin_Marine_商业介绍_中文版.docx' if zh else 'TailFin_Marine_Company_Profile_EN.docx'
    doc.save(out)
    print('saved', out)


# ---------------- CONTENT ----------------
TXT = {
 'zh': {
  'sub': '为硬核钓客而生，全球适用。',
  'lede': 'TailFin Marine 重新定义了"钓后处理"——将去鳞、剖切、真空锁鲜与离网供电整合进一套可车载、可登船、可入户的模块化系统。从皮卡尾门到远洋甲板，从冰湖帐篷到近海码头，让每一位钓客都能"随钓、随处理、随封存"。',
  'sections': [
   {'h':'执行摘要','body':[
     ('p','全球休闲垂钓正经历一场结构性增长：参与者数量连创新高，消费从"装备入门"走向"体验升级"，而"钓后处理（清洁与烹饪）"这一长期被忽视的环节，正在成为品类创新的高地。现有市场以被动式折叠去鳞台（HDPE 台面+水龙头+刀槽）为主，用户仍需自行携带并手动操作去鳞机、鱼刀与真空机——功能割裂、收纳繁琐、无法离网。'),
     ('p','TailFin Marine 创立的全新品类是"动力模块化钓后处理系统"：以统一基础平台+中央仓双槽叠层为核心，去鳞机、电动鱼刀、真空密封、循环水与离网电池即插即用。用户一次购买基础款，后续按钓法与预算按需叠加模块升级，形成可持续的复购与增值闭环。'),
     ('callout','核心主张 · 随钓、随净、随烹，随处皆可。一套系统，把"把鱼带回家"变成"在现场就把鱼处理好、封存好"。'),
   ]},
   {'h':'市场机会','body':[
     ('p','休闲垂钓已从区域性的小众爱好，成长为全球性的户外生活方式与情绪消费赛道。以下为 2025–2026 年公开资料中的关键数据：'),
     ('table',['区域/市场','规模与增速','关键信号'],[
       ['全球装备市场','休闲垂钓装备 2025 年约 148 亿美元，预计 2034 年达 244 亿美元（CAGR 5.7%）；广义渔具市场 2025 年约 165 亿美元，2032 年达 221 亿美元（CAGR 4.25%）。','轻量化、便携化、模块化套装备受青睐。'],
       ['美国','2024 年 5,790 万人参与（占人口 19%，历史新高）；新增首次钓鱼者 510 万；女性 2,130 万、咸水钓鱼 1,510 万均创新高。','家庭与年轻群体驱动，社媒钓鱼内容爆发。'],
       ['欧洲','休闲垂钓装备 2025 年约 22 亿美元（占全球约 15%，CAGR 5.0%）；海洋休闲渔业约 870 万海钓者，年支出 59 亿欧元。','环保法规趋严，"可持续垂钓"成主流。'],
       ['大洋洲','仅西澳一地即有 75 万休闲钓客，贡献 24 亿澳元经济。','"钓鱼+文旅+社区"模式成熟。'],
       ['中国','约 1.4 亿活跃钓友，钓具产业规模超 500 亿元；轻量化装备 2025–2029 CAGR 高达 17.3%，2029 年有望破 254 亿元。','年轻化、轻垂钓兴起，电商高速增长。'],
     ]),
     ('muted','趋势共识：①装备向轻量、便携、模块化演进；②智能与电动化渗透率提升；③环保与可持续成为购买要素；④"钓鱼+露营+文旅"融合，催生"钓后处理"新消费场景。'),
   ]},
   {'h':'品牌定位与使命','body':[
     ('p','我们刻意去除任何区域性标签——TailFin Marine 不为某一片大陆而生，而为全世界的硬核钓客设计。无论在北美的湖泊、欧洲的河道、澳洲的海岸，还是亚洲的野塘，钓后处理的核心痛点完全一致：鱼要尽快处理、要干净、要锁鲜、要能离网完成。'),
     ('h3','全球适用'),('muted','统一电压与插头适配方案，按全球主要市场认证路径设计，不绑定单一区域使用场景。'),
     ('h3','硬核级'),('muted','316 海军级不锈钢、食品级 HDPE、防飞溅结构，经盐雾与户外工况验证，对标专业钓客的耐用要求。'),
     ('h3','模块化'),('muted','一套基础平台，按需叠加去鳞/剖切/真空/供电模块，陪伴钓客从入门到专业。'),
     ('muted','使命：让"把今天的渔获变成今晚的餐桌"在任何地方都成为可能——干净、快速、可持续。'),
   ]},
   {'h':'关于 TailFin Marine','body':[
     ('p','TailFin Marine 是一家聚焦"钓后处理"赛道的户外装备品牌。我们由一群痴迷钓鱼的工程师与产品人创立，发现市面上的去鳞台只是"一张桌子"，而真正的痛点——动力处理、锁鲜封存、离网供电——长期处于割裂状态。于是我们以"中央仓模块化"为底层架构，把整套钓后处理流程压缩进一个可搬可装的系统里。'),
     ('p','我们采用研产销一体+供应链复用模式，核心结构件与珠三角成熟钓具/锂电/真空供应链深度协同，在保持专业级品质的同时具备极具竞争力的交付成本，便于在全球市场快速铺货与本地化适配。'),
     ('callout','我们不是在卖"一张更好的桌子"，而是在定义一个新品类——这正是分销商抢占空白货架、建立差异化的机会。'),
   ]},
   {'h':'产品体系','body':[
     ('p','全系围绕统一基础平台构建，价格梯度清晰，覆盖从入门到专业的完整需求，并为配件与耗材预留持续复购空间。'),
     ('table',['型号','定位','价格','要点'],[
       ['口袋去鳞机 Mini','入门款','$149','316 不锈钢去鳞头；12V 2200mAh 锂电 Type-C 充电；配鱼剪+10L 水袋+牛津布包。'],
       ['尾门工作站 Standard','基础款','$499','箱体 80×44×26cm；12V 100W 防水去鳞机；5L/min 循环水；HDPE 砧板；工具抽屉。'],
       ['移动鱼类工厂 Pro','旗舰款','$999','24V 200W 无刷去鳞机+电动鱼刀；真空≥-75kPa；24V 电池离网≥2h；BMS 双充。'],
       ['Pro 升级套装 Kit','升级套装','$499','真空模块+24V 电池+无刷去鳞机+电动鱼刀，Standard 一键升级至 Pro，立省 $57。'],
     ]),
     ('h3','配件与耗材（12 款，独立复购）'),
     ('muted','去鳞刀头 $29 · 防穿刺真空袋(20片) $19 · PU 进水管 $25 · 电动鱼刀 $79 · 鱼剪刀 $15 · 12V/24V 去鳞机 $89/$129 · 电池包 $149 · 真空模块 $199 · HDPE 砧板 $35 · 牛津布包 $25 · 折叠水袋 $18 等。全系通用，构成持续消耗与增值收入。'),
   ]},
   {'h':'模块化架构与核心技术','body':[
     ('ul',[
       '中央仓双槽叠层（盲插对位）：Standard 加购 4 模块即叠成 Pro，供电/信号/结构一次对位，热插拔不接线。',
       '316 海军级不锈钢：去鳞头与机身通过 72 小时盐雾测试，耐腐蚀、利落不伤肉。',
       '干湿两用真空密封 ≥-75kPa：隐藏式液体收集盒护泵，配防穿刺袋锁鲜防串味。',
       '离网供电 24V 6S1P：电池管理系统双充（110V+12V），单次≥2 小时，支持热插拔。',
       '循环水+防飞溅：PC 防飞溅罩+5L/min 循环给水，干净不脏手、不浪费水。',
       '统一数据层：价格/促销/上下架由后台单一数据源管理，前后端实时同步，便于本地化运营。',
     ]),
   ]},
   {'h':'应用场景','body':[
     ('h3','陆地 · 皮卡尾门/营地'),('muted','停车展开箱体，直接在尾门完成去鳞剖切与真空封存，无需回程处理、不脏车内。'),
     ('h3','冰湖 · 冰钓帐篷'),('muted','-20℃ 环境由电池独立供电，钓上即处理，热食与封存一气呵成。'),
     ('h3','近海 · 船甲板/码头'),('muted','登船即装，12V 车载供电持续作业；回港前已封装完毕，直达餐桌。'),
     ('h3','野塘 · 溪流轻垂钓'),('muted','Mini 口袋款+折叠水袋，轻装即兴出行，处理完随手收纳。'),
   ]},
   {'h':'竞品对比','body':[
     ('p','现有市场以被动式 HDPE 去鳞/剖鱼台为主——本质是"一块带水槽或刀槽的板"，动力去鳞、真空锁鲜与离网供电全部缺席。TailFin 以一套系统整合全流程，形成品类级差异。'),
     ('table',['能力维度','TailFin Standard/Pro','Goplus 折叠台','MAGMA Bait/Filet','RAILBLAZA/传统 HDPE'],[
       ['产品形态','动力模块化钓后处理系统','折叠水槽台','HDPE 剖鱼砧板','船用剖鱼砧板'],
       ['动力去鳞','✓ 12V/24V 电动去鳞机','✗ 无','✗ 无','✗ 无'],
       ['电动剖切','✓ 往复式电动鱼刀(Pro)','✗ 手动','✗ 手动','✗ 手动'],
       ['真空锁鲜','✓ 中央仓真空 ≥-75kPa','✗ 无','✗ 无','✗ 无'],
       ['离网供电','✓ 24V 电池 ≥2h+双充','✗ 需接水/无电','✗ 无','✗ 无'],
       ['循环/供水','✓ 5L/min 循环+防飞溅','△ 接花园水管','✗ 无','✗ 无'],
       ['模块化升级','✓ 基础→Pro 即插即用','✗ 单一形态','✗ 单一形态','✗ 单一形态'],
       ['材质','316 不锈钢+HDPE+ABS/PC','HDPE+喷粉钢腿','UV 稳定 HDPE','HDPE'],
       ['参考零售价','$499/$999（整机）','约 $93','约 $66–$400','约 $95–$120'],
       ['复购生态','✓ 12 配件+升级套装','△ 极少','△ 仅配安装座','△ 仅配安装座'],
     ]),
     ('callout','结论：竞品是"更好的台面"，TailFin 是"完整的钓后处理系统"。价格对标的是"去鳞台+电动去鳞机+电动鱼刀+真空机+便携电源"的组合总价，且一次购齐、即插即用、可离网——不是同一维度的竞争。'),
   ]},
   {'h':'为什么选择我们 · 分销合作价值','body':[
     ('ul',[
       '品类空白：市场尚无"动力模块化钓后处理"竞品，以差异化产品占据空白货架与心智。',
       '高复购：12 款配件/耗材+升级套装形成持续消耗收入，客户生命周期价值显著高于单品。',
       '价格梯度：$149–$999 全梯度覆盖，入门引流、旗舰树形象、升级套装提客单。',
       '区域保护：按市场划分代理与经销层级，保障合作伙伴利益。',
       '营销素材：提供多语言产品图、场景图、视频与电商详情页模板，降低本地化门槛。',
       '柔性供应：珠三角成熟供应链支撑小批量试单与快速补货，库存风险可控。',
       '认证支持：按目标市场（CE/FCC/RCM 等）提供合规与认证路径协助。',
     ]),
     ('callout','面向客户的价值同样清晰：省时、干净、锁鲜、离网——把"处理鱼"从负担变成钓鱼乐趣的延伸。'),
   ]},
   {'h':'分销政策与毛利测算','body':[
     ('p','我们采用清晰的三级价格体系，价格随层级与承诺量递进，保障渠道利润与区域秩序。下表为示意折扣与毛利模型，最终以正式报价单与合同为准。'),
     ('table',['渠道层级','拿货价(占MSRP)','渠道毛利率','首单起订量','年度返点','权益'],[
       ['授权零售商','70% MSRP','约 30%','10 台/混装','—','官方素材+培训'],
       ['区域经销商','60% MSRP','约 40%','50 台','销售额 2–4%','次区域独家+优先补货'],
       ['国家分销商','50% MSRP','约 50%','200 台/整柜','销售额 3–6%','区域独家+认证协助+联合营销基金'],
     ]),
     ('h3','毛利测算示例（国家分销商 50% MSRP 拿货）'),
     ('table',['型号','MSRP 建议零售','分销拿货价','单台毛利','毛利率'],[
       ['口袋去鳞机 Mini','$149','$74.5','$74.5','50%'],
       ['尾门工作站 Standard','$499','$249.5','$249.5','50%'],
       ['移动鱼类工厂 Pro','$999','$499.5','$499.5','50%'],
       ['Pro 升级套装 Kit','$499','$249.5','$249.5','50%'],
       ['配件/耗材（12 款）','$15–$199','约 45% MSRP','—','约 55%'],
     ]),
     ('callout','典型整柜（200 台）样例：100×Standard+60×Pro+40×Kit，MSRP 合计约 $149,760，分销拿货约 $74,880，渠道潜在毛利约 $74,880——尚未计入高毛利配件耗材的持续复购收入。'),
     ('ul',[
       '价格保护：统一 MSRP 与 MAP（最低广告价）政策，防止渠道价格战。',
       '区域独家：分销商在约定区域享独家代理权，设年度目标与保底采购量。',
       '账期与首单：首单预付、验厂后可谈账期；提供小批量试单降低门槛。',
       '联合营销基金(MDF)：达标分销商可获销售额一定比例的市场推广支持。',
       '认证与售后：按目标市场协助合规；提供备件包与保修政策支持本地售后。',
     ]),
     ('muted','注：以上折扣率、MOQ、返点比例均为示意区间，实际条款依市场、承诺量与合作深度一事一议。'),
   ]},
   {'h':'合作洽谈与数据来源','body':[
     ('p','TailFin Marine 期待与全球分销商、户外零售连锁、钓具专营店及文旅渠道建立合作。欢迎来信获取报价单、经销政策与样品方案。'),
     ('muted','数据来源（公开资料 2025–2026）：Dimension Market Research；PMarketResearch；RBFF & Outdoor Foundation；European Anglers Alliance/UEA；FRDC/Recfishwest；新华社/尚普咨询/头豹研究院/弗若斯特沙利文；竞品公开报价（Goplus 约 $93、MAGMA 约 $66–$400、RAILBLAZA 约 $95–$120）。市场规模引用公开第三方研究，产品参数与价格为当前目录口径，最终以正式报价单与合同为准。'),
   ]},
  ],
 },
 'en': {
  'sub': 'Engineered for serious anglers, everywhere.',
  'lede': 'TailFin Marine redefines post-catch processing — integrating scaling, filleting, vacuum sealing and off-grid power into a single modular system that rides in your truck, mounts on your deck, and stores at home. From tailgate to offshore deck, from ice-fishing tent to coastal pier, every angler can catch, process and preserve right on the spot.',
  'sections': [
   {'h':'Executive Summary','body':[
     ('p','Global recreational fishing is undergoing structural growth: participation is hitting record highs, and spending is shifting from entry-level gear to experience upgrades. Yet the long-overlooked "Clean & Cook" stage is emerging as the next frontier of category innovation. Today\u2019s market is dominated by passive folding fillet tables (HDPE surface + faucet + knife slots), forcing users to carry and manually operate separate scalers, knives and vacuum sealers.'),
     ('p','TailFin Marine has created an entirely new category: the Powered, Modular Catch & Cook System. Built on a unified base platform with a central-bay dual-slot stack, the scaler, electric fillet knife, vacuum sealer, water recirculation and off-grid battery are all plug-and-play. Customers buy the base unit once, then stack modular upgrades as their fishing style and budget evolve.'),
     ('callout','Core promise: Catch. Clean. Cook. Anywhere. One system that turns "bringing the fish home" into "processing and preserving it right where you caught it."'),
   ]},
   {'h':'Market Opportunity','body':[
     ('p','Recreational fishing has grown from a regional niche into a global outdoor lifestyle and experiential-spending category. Key figures from public 2025\u20132026 sources:'),
     ('table',['Region/Market','Size & Growth','Key Signals'],[
       ['Global Equipment','Recreational fishing equipment ~$14.8B in 2025, reaching $24.4B by 2034 (CAGR 5.7%); broader tackle market ~$16.5B in 2025, $22.1B by 2032 (CAGR 4.25%).','Lightweight, portable, modular kits favored.'],
       ['United States','57.9M participants in 2024 (19% of population, record); 5.1M first-timers; 21.3M women and 15.1M saltwater anglers, both records.','Families and younger demographics; social-media boom.'],
       ['Europe','Equipment ~$2.2B in 2025 (~15% of global, CAGR 5.0%); marine recreational fishing ~8.7M sea anglers, \u20ac5.9B annual spend.','Tightening regulation; sustainability mainstream.'],
       ['Oceania','Western Australia alone: 750,000 anglers, A$2.4B economy.','Mature fishing + tourism + community model.'],
       ['China','~140M active anglers; tackle industry > RMB 50B; lightweight gear CAGR 17.3% (2025\u20132029), > RMB 25.4B by 2029.','Younger base, light-angling rise, e-commerce growth.'],
     ]),
     ('muted','Consensus trends: (1) gear evolving toward lightweight/portable/modular; (2) rising smart & electric penetration; (3) eco & sustainability as purchase drivers; (4) fishing + camping + tourism fusion creating a new post-catch processing scenario.'),
   ]},
   {'h':'Positioning & Mission','body':[
     ('p','We deliberately removed any regional label \u2014 TailFin Marine is designed for serious anglers worldwide. Whether on North American lakes, European rivers, Australian coasts or Asian ponds, the pain points of post-catch processing are identical: fish must be processed fast, kept clean, sealed fresh, and handled off-grid.'),
     ('h3','Borderless'),('muted','Unified voltage/plug-adapter scheme; designed for major-market certification paths, not tied to any single region.'),
     ('h3','Built Tough'),('muted','316 marine-grade stainless steel, food-grade HDPE, splash-proof construction; validated by salt-spray and outdoor testing.'),
     ('h3','Modular by Design'),('muted','One base platform; add scaling/filleting/vacuum/power modules on demand \u2014 growing from entry to pro.'),
     ('muted','Mission: make "turning today\u2019s catch into tonight\u2019s dinner" possible anywhere \u2014 clean, fast, sustainable.'),
   ]},
   {'h':'About TailFin Marine','body':[
     ('p','TailFin Marine is an outdoor-gear brand focused on the post-catch processing category. Founded by fishing-obsessed engineers and product designers, we saw that market fillet tables are just "a table," while the real pain points \u2014 powered processing, freshness sealing, off-grid power \u2014 remained fragmented. So we built a central-bay modular architecture that compresses the entire Catch & Cook workflow into one portable system.'),
     ('p','We operate on an integrated R&D\u2013manufacturing\u2013sales model with supply-chain reuse: core parts leverage the mature tackle/lithium-battery/vacuum supply chains of the Pearl River Delta, delivering professional-grade quality at competitive cost.'),
     ('callout','We are not selling "a better table" \u2014 we are defining a new category. That is the opportunity for distributors to claim an empty shelf and build differentiation.'),
   ]},
   {'h':'Product System','body':[
     ('p','The entire line is built on a unified base platform with a clear price ladder, from entry to professional, with ample room for recurring accessory and consumable purchases.'),
     ('table',['Model','Tier','Price','Highlights'],[
       ['Pocket Scaler Mini','LITE','$149','316 stainless head; 12V 2200mAh Type-C battery; fish scissors + 10L water bag + Oxford bag.'],
       ['Tailgate Workstation Standard','STANDARD','$499','Enclosure 80\u00d744\u00d726cm; 12V 100W waterproof scaler; 5L/min recirc; HDPE board; tool drawer.'],
       ['Mobile Fish Factory Pro','PRO','$999','24V 200W brushless scaler + electric knife; vacuum \u2265-75kPa; 24V battery \u22652h; BMS dual charge.'],
       ['Pro Upgrade Kit','UPGRADE','$499','Vacuum module + 24V battery + brushless scaler + electric knife; upgrade Standard to Pro, save $57.'],
     ]),
     ('h3','Accessories & Consumables (12 SKUs, repeat purchase)'),
     ('muted','Scaling head $29 · Vacuum bags 20-pack $19 · PU hose $25 · Electric knife $79 · Fish scissors $15 · 12V/24V scaler $89/$129 · Battery pack $149 · Vacuum module $199 · HDPE board $35 · Oxford bag $25 · Water bag $18, etc. Cross-compatible across the line \u2014 recurring revenue.'),
   ]},
   {'h':'Architecture & Core Tech','body':[
     ('ul',[
       'Central-bay dual-slot stack (blind-mate): add 4 modules to turn Standard into Pro; power/signal/structure align in one motion, hot-swappable.',
       '316 marine-grade stainless steel: head and body pass 72-hour salt-spray test.',
       'Wet/dry vacuum sealing \u2265-75kPa: hidden liquid-collection box protects the pump; puncture-resistant bags lock freshness.',
       'Off-grid power 24V 6S1P: BMS dual charging (110V AC + 12V vehicle), \u22652 hours per charge, hot-swappable.',
       'Recirculating water + splash guard: PC shield + 5L/min supply \u2014 clean hands, no wasted water.',
       'Unified data layer: pricing/promotions/availability from a single back-end source, synced in real time.',
     ]),
   ]},
   {'h':'Use Cases','body':[
     ('h3','Land · Truck Tailgate / Campsite'),('muted','Park, unfold, and scale/fillet/vacuum-seal right on the tailgate \u2014 no home processing, no mess in the vehicle.'),
     ('h3','Ice Lake · Ice-Fishing Tent'),('muted','In -20\u2103 the battery powers everything independently; process as you catch.'),
     ('h3','Coastal · Boat Deck / Pier'),('muted','Mount on board, run on 12V vehicle power; everything sealed before returning to port.'),
     ('h3','Freshwater · Streams & Ponds'),('muted','Mini pocket unit + folding water bag for light, spontaneous trips.'),
   ]},
   {'h':'Competitive Comparison','body':[
     ('p','Today\u2019s market is dominated by passive HDPE fillet/cleaning tables \u2014 essentially "a board with a sink or knife slots." Powered scaling, vacuum sealing and off-grid power are all absent. TailFin integrates the entire workflow into one system.'),
     ('table',['Capability','TailFin Standard/Pro','Goplus Table','MAGMA Bait/Filet','RAILBLAZA/Classic HDPE'],[
       ['Form factor','Powered modular system','Folding sink table','HDPE fillet board','Marine fillet board'],
       ['Powered scaling','\u2713 12V/24V scaler','\u2717 None','\u2717 None','\u2717 None'],
       ['Electric filleting','\u2713 Reciprocating knife (Pro)','\u2717 Manual','\u2717 Manual','\u2717 Manual'],
       ['Vacuum sealing','\u2713 Central-bay \u2265-75kPa','\u2717 None','\u2717 None','\u2717 None'],
       ['Off-grid power','\u2713 24V battery \u22652h','\u2717 Mains/no power','\u2717 None','\u2717 None'],
       ['Water/recirc','\u2713 5L/min + splash guard','\u25b3 Garden hose','\u2717 None','\u2717 None'],
       ['Modular upgrade','\u2713 Base\u2192Pro plug-and-play','\u2717 Single form','\u2717 Single form','\u2717 Single form'],
       ['Materials','316 stainless + HDPE + ABS/PC','HDPE + steel legs','UV HDPE','HDPE'],
       ['Reference price','$499/$999 (system)','~$93','~$66\u2013$400','~$95\u2013$120'],
       ['Repeat ecosystem','\u2713 12 accessories + kit','\u25b3 Minimal','\u25b3 Mounts only','\u25b3 Mounts only'],
     ]),
     ('callout','Bottom line: competitors sell "a better table"; TailFin delivers "a complete post-catch processing system," benchmarked against the combined cost of table + scaler + knife + sealer + portable power \u2014 bought once, plug-and-play, off-grid. Not the same category of competition.'),
   ]},
   {'h':'Why Partner With Us','body':[
     ('ul',[
       'Empty category: no powered modular post-catch competitor exists yet \u2014 claim an empty shelf and mindshare.',
       'High repeat: 12 accessories + upgrade kit create recurring revenue; LTV far exceeds single-item sales.',
       'Price ladder: full $149\u2013$999 coverage \u2014 entry drives traffic, flagship builds image, kit lifts basket.',
       'Territory protection: agent/dealer tiers segmented by market.',
       'Marketing assets: multilingual imagery, video and e-commerce templates.',
       'Flexible supply: mature supply chain supports small trial orders and fast replenishment.',
       'Certification support: assistance with CE/FCC/RCM paths per target market.',
     ]),
     ('callout','Value to end customers: time-saving, clean, freshness-locked, off-grid \u2014 turning "processing fish" into an extension of the fun of fishing.'),
   ]},
   {'h':'Dealer Policy & Margin Model','body':[
     ('p','We use a clear three-tier pricing structure, stepping down as tier and commitment volume increase. The table below is an illustrative discount and margin model; final terms are governed by the formal quotation and contract.'),
     ('table',['Channel Tier','Buy Price','Margin','First-Order MOQ','Annual Rebate','Benefits'],[
       ['Authorized Retailer','70% MSRP','~30%','10 units/mixed','\u2014','Assets + training'],
       ['Regional Dealer','60% MSRP','~40%','50 units','2\u20134% of sales','Sub-region exclusivity + priority restock'],
       ['Country Distributor','50% MSRP','~50%','200 units/container','3\u20136% of sales','Territory exclusivity + certification + co-marketing fund'],
     ]),
     ('h3','Margin Example (Country Distributor at 50% MSRP)'),
     ('table',['Model','MSRP','Distributor Price','Margin/Unit','Margin %'],[
       ['Pocket Scaler Mini','$149','$74.5','$74.5','50%'],
       ['Tailgate Workstation Standard','$499','$249.5','$249.5','50%'],
       ['Mobile Fish Factory Pro','$999','$499.5','$499.5','50%'],
       ['Pro Upgrade Kit','$499','$249.5','$249.5','50%'],
       ['Accessories/Consumables (12)','$15\u2013$199','~45% MSRP','\u2014','~55%'],
     ]),
     ('callout','Sample container (200 units): 100\u00d7Standard + 60\u00d7Pro + 40\u00d7Kit \u2014 total MSRP ~$149,760, distributor cost ~$74,880, potential channel margin ~$74,880 \u2014 before recurring accessory revenue.'),
     ('ul',[
       'Price protection: unified MSRP and MAP policy prevents price wars.',
       'Territory exclusivity: exclusive rights within an agreed territory, with annual targets and minimum commitments.',
       'Terms & first order: first order prepaid; credit negotiable after factory audit; trial orders available.',
       'Market Development Fund (MDF): qualifying distributors receive a % of sales toward marketing.',
       'Certification & after-sales: compliance assistance per market; spare-parts kits and warranty support.',
     ]),
     ('muted','Note: discount rates, MOQs and rebates above are illustrative ranges; actual terms are case-by-case based on market, volume and partnership depth.'),
   ]},
   {'h':'Partnership & Data Sources','body':[
     ('p','TailFin Marine welcomes partnerships with global distributors, outdoor retail chains, tackle specialty stores and tourism channels. Contact us for quotations, dealer policy and sample programs.'),
     ('muted','Data sources (public, 2025\u20132026): Dimension Market Research; PMarketResearch; RBFF & Outdoor Foundation; European Anglers Alliance/UEA; FRDC/Recfishwest; Xinhua/Sunpu/LeadLeo/Frost & Sullivan; competitor public pricing (Goplus ~$93, MAGMA ~$66\u2013$400, RAILBLAZA ~$95\u2013$120). Market figures cite public third-party research; specs and prices reflect the current catalog; final terms governed by the formal quotation and contract.'),
   ]},
  ],
 },
}

if __name__ == '__main__':
    build('zh')
    build('en')
