# -*- coding: utf-8 -*-
"""Generic HTML -> .docx converter for TailFin Marine docs.

Keeps <table> as native Word tables (with header row + grid style),
renders headings / paragraphs / lists / blockquotes as Word elements,
and preserves inline bold/italic. No HTML content is modified.
"""
import sys
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0b, 0x25, 0x45)
TEAL = RGBColor(0x1b, 0x99, 0x8b)
INK = RGBColor(0x1f, 0x29, 0x33)
GRAY = RGBColor(0x5b, 0x6b, 0x7b)
TEALBG = "D9F0EB"
GRAYBG = "F2F4F7"
CJK_FONT = "Microsoft YaHei"


def set_cjk(run, font=CJK_FONT):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_inline(paragraph, element):
    """Append inline text (with b/i) from an element into a paragraph."""
    if element is None:
        return
    if isinstance(element, str):
        if element.strip():
            paragraph.add_run(element)
        return
    if element.name is None:
        txt = element.string
        if txt:
            paragraph.add_run(txt)
        return
    if element.name in ('b', 'strong'):
        run = paragraph.add_run(element.get_text())
        run.bold = True
        set_cjk(run)
    elif element.name in ('i', 'em'):
        run = paragraph.add_run(element.get_text())
        run.italic = True
        set_cjk(run)
    elif element.name == 'br':
        paragraph.add_run('\n')
    else:
        for child in element.children:
            add_inline(paragraph, child)


def add_table(doc, table_el):
    rows = table_el.find_all('tr', recursive=False)
    if not rows:
        rows = table_el.find_all('tr')
    if not rows:
        return
    ncols = 0
    for r in rows:
        ncols = max(ncols, len(r.find_all(['td', 'th'], recursive=False) or r.find_all(['td', 'th'])))
    if ncols == 0:
        return
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = 'Table Grid'
    for ri, r in enumerate(rows):
        cells = r.find_all(['td', 'th'], recursive=False) or r.find_all(['td', 'th'])
        is_header = (r.find('th') is not None) or (cells and cells[0].name == 'th')
        for ci in range(ncols):
            cell = t.cell(ri, ci)
            src = cells[ci] if ci < len(cells) else None
            if src is not None:
                text = src.get_text('\n', strip=True)
                para = cell.paragraphs[0]
                run = para.add_run(text)
                run.font.size = Pt(9)
                if is_header:
                    run.bold = True
                    run.font.color.rgb = NAVY
                    set_cell_bg(cell, TEALBG)
                else:
                    run.font.color.rgb = INK
                    if ri % 2 == 1:
                        set_cell_bg(cell, GRAYBG)
                set_cjk(run)
    doc.add_paragraph()


def block_handler(doc, element):
    """Recursively render a block element into the docx."""
    if element is None:
        return
    if isinstance(element, str):
        return
    if element.name is None:
        return

    name = element.name
    if name == 'table':
        add_table(doc, element)
        return
    if name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        sizes = {'h1': 20, 'h2': 16, 'h3': 13.5, 'h4': 12, 'h5': 11, 'h6': 10.5}
        cls = element.get('class')
        cls = ' '.join(cls) if cls else ''
        p = doc.add_paragraph()
        run = p.add_run(element.get_text(strip=True))
        run.bold = True
        run.font.size = Pt(sizes.get(name, 12))
        run.font.color.rgb = TEAL if name in ('h3', 'h4') else NAVY
        if 'lede' in cls:
            run.font.color.rgb = NAVY
        set_cjk(run)
        if name == 'h2':
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '1B998B')
            pbdr.append(bottom); pPr.append(pbdr)
        return
    if name == 'p':
        cls = element.get('class')
        cls = ' '.join(cls) if cls else ''
        muted = 'muted' in cls
        p = doc.add_paragraph()
        pcls = None
        if 'lede' in cls:
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), NAVY_HEX)
            pPr.append(shd)
        if 'en' in cls:
            pass
        for child in element.children:
            add_inline(p, child)
        for run in p.runs:
            run.font.size = Pt(13 if 'lede' in cls else 10.5)
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff) if 'lede' in cls else (GRAY if muted else INK)
            set_cjk(run)
        return
    if name in ('ul', 'ol'):
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, li)
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = INK
                set_cjk(run)
        return
    if name in ('blockquote',):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), TEALBG)
        pPr.append(shd)
        add_inline(p, element)
        for run in p.runs:
            run.font.size = Pt(10.5); run.bold = True
            run.font.color.rgb = TEAL; set_cjk(run)
        return
    if name in ('div', 'section', 'article', 'main', 'header', 'footer', 'body', 'span'):
        for child in element.children:
            block_handler(doc, child)
        return
    # fallback: recurse
    for child in element.children:
        block_handler(doc, child)


NAVY_HEX = "0B2545"


def convert(html_path, docx_path):
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = CJK_FONT
    style.element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)

    body = soup.body or soup
    for child in body.children:
        block_handler(doc, child)
    doc.save(docx_path)
    print('saved', docx_path)


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('usage: gen_html_docx.py <in.html> <out.docx>')
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
